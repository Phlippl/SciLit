"""
DOCX-Parser für SciLit
--------------------
Parser für Microsoft Word-Dokumente.
"""

import logging
from typing import Dict, Tuple, Any

# DOCX-Verarbeitung
try:
    import docx
except ImportError:
    docx = None

from app.core.document.parsers.base_parser import DocumentParser, DocumentParsingError

# Logger konfigurieren
logger = logging.getLogger("scilit.document.parsers.docx")

class DOCXParser(DocumentParser):
    """
    Parser für DOCX-Dokumente.
    
    Dieser Parser extrahiert Text und Metadaten aus DOCX-Dateien 
    mit dem python-docx Paket.
    """
    
    def parse(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parst ein DOCX-Dokument und extrahiert Text und Metadaten.
        
        Args:
            filepath: Pfad zum DOCX
            
        Returns:
            Tuple aus (extrahierter Text, Dictionary mit Metadaten)
            
        Raises:
            DocumentParsingError: Bei Problemen mit dem Parsing
        """
        if not docx:
            raise DocumentParsingError("python-docx ist nicht installiert")
        
        try:
            # Metadaten aus Dateiinformationen
            metadata = self._extract_basic_metadata(filepath)
            
            # Öffne das DOCX
            doc = docx.Document(filepath)
            
            # Text extrahieren
            text = []
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            
            # Metadaten aus DOCX-Eigenschaften
            core_properties = doc.core_properties
            
            if core_properties.title:
                metadata["title"] = core_properties.title
            
            if core_properties.author:
                metadata["author"] = [author.strip() for author in core_properties.author.split(';')]
            
            if core_properties.keywords:
                metadata["keywords"] = core_properties.keywords
            
            if core_properties.created:
                metadata["created"] = core_properties.created.isoformat()
                if core_properties.created.year:
                    metadata["year"] = core_properties.created.year
            
            if core_properties.modified:
                metadata["modified"] = core_properties.modified.isoformat()
            
            # Seitenzahl abschätzen (paragraph/40 als grobe Schätzung)
            self.page_count = max(1, len(doc.paragraphs) // 40)
            metadata["page_count"] = self.page_count
            
            return self._clean_text("\n\n".join(text)), metadata
            
        except Exception as e:
            raise DocumentParsingError(f"Fehler beim DOCX-Parsing: {str(e)}")